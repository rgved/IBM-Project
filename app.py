import streamlit as st
import json
import docx
import fitz  # PyMuPDF for PDF text extraction
from io import BytesIO
from model import companion_feedback, summarise_text  # functions from model.py
import re
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

st.set_page_config(page_title="AI Companion Tutor", layout="wide")

# ---------------------------
# File Parsing Helpers
# ---------------------------
def pdf_to_text(file):
    try:
        pdf_data = file.read()
        pdf_file = BytesIO(pdf_data)
        doc = fitz.open(stream=pdf_file, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ PDF extraction failed: {e}")
        return ""


def docx_to_text(file):
    try:
        doc_obj = docx.Document(file)
        return "\n".join([p.text for p in doc_obj.paragraphs if p.text.strip()]).strip()
    except Exception as e:
        st.error(f"❌ DOCX extraction failed: {e}")
        return ""


def detect_question(line):
    keywords = [
        r'\bdefine\b', r'\bdescribe\b', r'\bshow\b', r'\billustrate\b',
        r'\belaborate\b', r'\bexplain\b', r'\bgive\b',
        r'\bwho\b', r'\bwhat\b', r'\bwhere\b', r'\bwhen\b', r'\bwhy\b', r'\bhow\b'
    ]
    pattern = r'(\?$|:\s*$|' + "|".join(keywords) + r')'
    return re.search(pattern, line.strip(), flags=re.IGNORECASE)


def smart_parse_text_to_json(raw_text):
    raw_text = re.sub(r'\n+', '\n', raw_text.strip())
    lines = raw_text.split("\n")
    questions, current_q, current_a = [], None, []

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if detect_question(line) or re.match(r'^\d+[\).]', line):
            if current_q:
                questions.append({
                    "question_id": f"Q{len(questions)+1}",
                    "question": current_q,
                    "student_answer": " ".join(current_a).strip(),
                    "correct_answer": ""
                })
            current_q, current_a = line, []
        else:
            current_a.append(line)
    if current_q:
        questions.append({
            "question_id": f"Q{len(questions)+1}",
            "question": current_q,
            "student_answer": " ".join(current_a).strip(),
            "correct_answer": ""
        })
    return questions


# ---------------------------
# Export Helpers
# ---------------------------
def export_summary_as_pdf(summary_text):
    """Generate a downloadable PDF from summary text."""
    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    y = height - 50

    for line in summary_text.split('\n'):
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line)
        y -= 15
    c.save()
    pdf_buffer.seek(0)
    return pdf_buffer


def export_summary_as_docx(summary_text):
    """Generate a downloadable DOCX file from summary text."""
    doc = Document()
    doc.add_heading("AI Summary", level=1)
    for para in summary_text.split("\n"):
        doc.add_paragraph(para)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------
# Main UI
# ---------------------------
st.title("🤝 StuBard - AI Companion Tutor")
st.write("Your AI-powered learning partner for guided feedback or summaries!")

mode = st.radio("Choose Mode", ["🎓 Guidance Mode", "🧾 Summarizer Mode"])

# ===============================
# GUIDANCE MODE
# ===============================
if mode == "🎓 Guidance Mode":
    st.subheader("✏️ Guidance Mode - Get feedback on your answers")
    upload_option = st.radio("Choose Input Method", ["Manual Input", "Upload File"])

    question = ""
    student_answer = ""
    correct_answer = ""

    # --- Manual Input ---
    if upload_option == "Manual Input":
        question = st.text_area("Enter your Question")
        student_answer = st.text_area("Enter your Answer")
        correct_answer = st.text_area("Correct Answer (optional, for better guidance)")

    # --- File Upload ---
    elif upload_option == "Upload File":
        file = st.file_uploader("Upload a JSON, PDF, or DOCX file", type=["json", "pdf", "docx"])
        if file:
            st.success(f"✅ Uploaded: {file.name}")
            if file.name.endswith(".pdf"):
                raw_text = pdf_to_text(file)
                parsed = smart_parse_text_to_json(raw_text)
            elif file.name.endswith(".docx"):
                raw_text = docx_to_text(file)
                parsed = smart_parse_text_to_json(raw_text)
            elif file.name.endswith(".json"):
                parsed = json.load(file)
            else:
                parsed = []

            if parsed:
                q_idx = st.number_input("Select Question Index", min_value=1, max_value=len(parsed), value=1)
                selected = parsed[q_idx - 1]
                question = selected.get("question", "")
                student_answer = selected.get("student_answer", "")
                correct_answer = selected.get("correct_answer", "")

                st.write(f"**Question:** {question}")
                st.write(f"**Student Answer:** {student_answer}")

    max_score = st.number_input("Max Score", min_value=1, max_value=10, value=5)

    if st.button("Get Guidance"):
        if not question or not student_answer:
            st.warning("Please provide a question and your answer.")
        else:
            with st.spinner("Generating feedback..."):
                result = companion_feedback(question, student_answer, correct_answer, max_score)

            st.subheader("📢 Feedback")
            st.write(result.get("feedback", "No feedback available"))

            st.subheader("🔑 Keywords for a Perfect Answer")
            keywords = result.get("keywords", [])
            st.write(", ".join(keywords) if keywords else "No keywords found")

            st.subheader("🚀 Steps to Improve")
            steps = result.get("improvement_steps", [])
            if steps:
                for step in steps:
                    st.markdown(f"- {step}")
            else:
                st.write("No improvement steps available")

# ===============================
# SUMMARIZER MODE
# ===============================
elif mode == "🧾 Summarizer Mode":
    st.subheader("🧠 Summarizer Mode - Get concise summaries of long texts or documents")

    summ_option = st.radio("Choose Input Method", ["Manual Input", "Upload File"])
    user_text = ""

    if summ_option == "Manual Input":
        user_text = st.text_area("Paste or type the text you want to summarize")

    elif summ_option == "Upload File":
        file = st.file_uploader("Upload a PDF or DOCX file to summarize", type=["pdf", "docx"])
        if file:
            st.success(f"✅ Uploaded: {file.name}")
            if file.name.endswith(".pdf"):
                user_text = pdf_to_text(file)
            elif file.name.endswith(".docx"):
                user_text = docx_to_text(file)

    if st.button("Summarize"):
        if not user_text.strip():
            st.warning("Please enter or upload text to summarize.")
        else:
            with st.spinner("Generating summary..."):
                summary = summarise_text(user_text)

            st.subheader("🪄 Summary")
            st.write(summary if summary else "No summary generated.")

            # --- Export Buttons ---
            if summary:
                col1, col2 = st.columns(2)
                with col1:
                    pdf_buffer = export_summary_as_pdf(summary)
                    st.download_button(
                        label="📄 Download Summary as PDF",
                        data=pdf_buffer,
                        file_name="summary.pdf",
                        mime="application/pdf"
                    )
                with col2:
                    docx_buffer = export_summary_as_docx(summary)
                    st.download_button(
                        label="📝 Download Summary as DOCX",
                        data=docx_buffer,
                        file_name="summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
